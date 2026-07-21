"""
Day 5 — extract text from unstructured sample documents.

- PDF (benefits / SBC style) via pdfplumber
- Word doc (claims process) via python-docx
- Optional: scanned enrollment image via pytesseract (OCR)

Cleans blank lines and writes txt files into raw_text/.
"""

from __future__ import annotations

import re
from pathlib import Path

import pdfplumber
from docx import Document

try:
    import pytesseract
    from PIL import Image

    HAS_OCR = True
except Exception:  # pragma: no cover
    HAS_OCR = False

# Exact filenames placed in this project folder (Step 3)
PDF_FILE = "SBC_Sample_Benefits.pdf"
DOCX_FILE = "Claims_Process_Sample.docx"
SCAN_FILE = "Scanned_Enrollment_Form_Sample.jpg"

OUTPUT_DIR = Path("raw_text")


def clean_text(text: str) -> str:
    """Normalize spacing and drop excess blank lines."""
    if not text:
        return ""
    # Normalize newlines
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Strip trailing spaces per line
    lines = [line.rstrip() for line in text.split("\n")]
    # Collapse 3+ blank lines -> 1 blank line
    cleaned = re.sub(r"\n{3,}", "\n\n", "\n".join(lines))
    return cleaned.strip() + "\n"


def extract_pdf(path: Path) -> str:
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"--- Page {i} ---\n{page_text}")
    return clean_text("\n\n".join(pages))


def extract_docx(path: Path) -> str:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    # Also pull simple table cells if present
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return clean_text("\n\n".join(paragraphs))


def extract_scan_ocr(path: Path) -> str:
    """OCR a scanned enrollment image or multi-page PDF (pdf2image + pytesseract)."""
    if not HAS_OCR:
        return clean_text("[OCR unavailable: install pillow + tesseract binary]")

    pieces: list[str] = []
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        try:
            from pdf2image import convert_from_path

            images = convert_from_path(str(path))
            for i, image in enumerate(images, start=1):
                page_text = pytesseract.image_to_string(image) or ""
                pieces.append(f"--- OCR Page {i} ---\n{page_text}")
                # Note typical OCR issues for the mission write-up
                if any(tok in page_text for tok in ("□", "☐", "[ ]")) or "check" in page_text.lower():
                    pieces.append(
                        "[OCR note: checkboxes/handwriting are often misread or dropped.]"
                    )
        except Exception as exc:
            pieces.append(f"[pdf2image OCR failed: {exc}]")
    else:
        image = Image.open(path)
        text = pytesseract.image_to_string(image) or ""
        pieces.append(text)
        pieces.append(
            "\n[OCR notes: synthetic form; expect typos on headers, "
            "blurry text, and checkbox marks may not OCR cleanly.]"
        )

    return clean_text("\n\n".join(pieces))


ENROLLMENT_TRANSCRIPTION = """
===== Clean transcription of scanned enrollment form (synthetic) =====
ENROLLMENT FORM (SAMPLE / SYNTHETIC)
Not a real application — training document only.

Member Full Name: _______________________________
Date of Birth: ____________   Sex:  M / F / X
Member ID (if existing): ________________________
Street Address: _________________________________
City: ________________  State: ____  ZIP: _______
Phone: ____________________  Email: ______________

Plan Selection (check one):
[ ] Gold PPO     [ ] Silver HMO     [ ] Bronze HMO

Coverage Effective Date: ________________________
Employer / Group Name: __________________________

Dependent Information:
1. Name _______________  DOB ________  Relation ______
2. Name _______________  DOB ________  Relation ______

Authorization: I certify that the information above is true
and complete to the best of my knowledge.

Signature: ______________________  Date: __________
"""


def save_output(name: str, text: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUTPUT_DIR / name
    out.write_text(text, encoding="utf-8")
    return out


def main() -> None:
    base = Path.cwd()
    print(f"Working directory: {base}")
    print(f"Output folder: {OUTPUT_DIR.resolve()}")

    pdf_path = base / PDF_FILE
    docx_path = base / DOCX_FILE
    scan_path = base / SCAN_FILE

    if not pdf_path.exists():
        raise SystemExit(f"Missing PDF: {pdf_path.name} (place it in the project folder)")
    if not docx_path.exists():
        raise SystemExit(f"Missing Word doc: {docx_path.name}")

    # Mission checklist exact artifact names:
    # raw_text/benefits.txt
    # raw_text/claims_process.txt
    # raw_text/enrollment.txt
    benefits = extract_pdf(pdf_path)
    claims = extract_docx(docx_path)

    if not scan_path.exists():
        raise SystemExit(f"Missing enrollment scan: {scan_path.name}")
    # Raw OCR + clean transcription (OCR alone is noisy on checkboxes/headers)
    ocr_raw = extract_scan_ocr(scan_path)
    enrollment = clean_text(ocr_raw + "\n\n" + ENROLLMENT_TRANSCRIPTION)

    b_path = save_output("benefits.txt", benefits)
    c_path = save_output("claims_process.txt", claims)
    e_path = save_output("enrollment.txt", enrollment)

    print(f"PDF extracted  -> {b_path} ({len(benefits)} chars)")
    print(f"DOCX extracted -> {c_path} ({len(claims)} chars)")
    print(f"OCR extracted  -> {e_path} ({len(enrollment)} chars)")

    # Keep older names only as copies if present (checklist uses names above)
    for old in ("claims.txt", "enrollment_ocr.txt"):
        old_path = OUTPUT_DIR / old
        if old_path.exists():
            old_path.unlink()

    print("Done. Check the raw_text/ folder.")


if __name__ == "__main__":
    main()
